from collections import Counter
from difflib import SequenceMatcher
from io import BytesIO
import json
import re

from django.http import HttpResponse, JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt

try:
    from sentence_transformers import SentenceTransformer, util
except ImportError:
    SentenceTransformer = None
    util = None

try:
    import spacy
except ImportError:
    spacy = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    import docx
except ImportError:
    docx = None


semantic_model = None
nlp = None


def load_models():
    global semantic_model, nlp

    if semantic_model is None and SentenceTransformer is not None:
        try:
            semantic_model = SentenceTransformer("all-MiniLM-L6-v2")
        except Exception:
            semantic_model = None

    if nlp is None and spacy is not None:
        try:
            nlp = spacy.load("en_core_web_sm")
        except Exception:
            nlp = None


load_models()


SKILLS = {
    "programming": [
        "python", "java", "javascript", "c++", "c#", "php", "ruby", "go", "rust", "swift",
        "kotlin", "typescript", "scala", "perl", "r", "matlab"
    ],
    "web": [
        "html", "css", "react", "angular", "vue", "node", "express", "django", "flask",
        "spring", "laravel", "asp.net", "jquery", "bootstrap", "sass", "less"
    ],
    "database": [
        "sql", "mysql", "postgresql", "mongodb", "redis", "oracle", "sqlite",
        "cassandra", "elasticsearch", "firebase"
    ],
    "tools": [
        "git", "github", "gitlab", "docker", "kubernetes", "jenkins", "aws", "azure",
        "gcp", "linux", "bash", "powershell", "vscode", "intellij", "eclipse"
    ],
    "frameworks": [
        "django", "flask", "fastapi", "tensorflow", "pytorch", "scikit-learn",
        "pandas", "numpy", "opencv", "keras", "hadoop", "spark"
    ],
    "soft_skills": [
        "agile", "scrum", "kanban", "leadership", "communication", "teamwork",
        "problem solving", "critical thinking"
    ],
    "other": [
        "machine learning", "data science", "artificial intelligence", "api",
        "rest api", "graphql", "microservices", "testing", "ci/cd", "devops"
    ]
}

ALL_SKILLS = [skill for category in SKILLS.values() for skill in category]

SKILL_CATEGORIES = {
    "Programming": {"python", "java", "javascript", "c++", "c#", "php", "ruby", "go", "rust", "swift", "kotlin", "typescript", "scala", "perl", "r", "matlab"},
    "Web Development": {
        "html", "css", "react", "angular", "vue", "node", "express", "django", "flask",
        "spring", "laravel", "asp.net", "jquery", "bootstrap", "sass", "less", "api", "rest api"
    },
    "Database": {"sql", "mysql", "postgresql", "mongodb", "redis", "oracle", "sqlite", "cassandra", "elasticsearch", "firebase"},
    "Cloud": {"aws", "azure", "gcp", "docker", "kubernetes"},
    "Tools": {"git", "github", "gitlab", "jenkins", "linux", "bash", "powershell", "vscode", "intellij", "eclipse", "ci/cd"}
}

CRITICAL_SKILL_HINTS = [
    "must", "required", "mandatory", "strong", "expert", "proficient",
    "proficiency", "hands-on", "experience with", "knowledge of", "need", "seeking"
]

SECTION_ALIASES = {
    "summary": {
        "summary", "professional summary", "profile", "career objective", "objective", "about"
    },
    "skills": {
        "skills", "technical skills", "core skills", "key skills", "technologies", "technical proficiencies"
    },
    "experience": {
        "experience", "work experience", "professional experience", "employment history", "internship", "internships"
    },
    "projects": {
        "projects", "project", "academic projects", "personal projects", "project experience"
    },
    "education": {
        "education", "academic background", "academic qualification", "qualifications", "qualification"
    },
    "certifications": {
        "certifications", "certification", "licenses", "courses"
    }
}

JOB_ROLES = [
    {
        "title": "Frontend Developer",
        "description": "Build interactive user interfaces using modern web technologies.",
        "keySkills": ["html", "css", "javascript", "react", "vue"]
    },
    {
        "title": "Backend Developer",
        "description": "Design and implement server-side applications and APIs.",
        "keySkills": ["python", "django", "node", "sql", "api"]
    },
    {
        "title": "Full Stack Developer",
        "description": "Work across frontend and backend to build complete web applications.",
        "keySkills": ["python", "django", "javascript", "html", "css"]
    },
    {
        "title": "Data Analyst",
        "description": "Analyze datasets to derive business insights.",
        "keySkills": ["python", "sql", "excel", "data science"]
    },
    {
        "title": "DevOps Engineer",
        "description": "Manage CI/CD pipelines and infrastructure automation.",
        "keySkills": ["docker", "kubernetes", "aws", "git", "ci/cd"]
    },
    {
        "title": "Machine Learning Engineer",
        "description": "Develop and deploy ML models to production.",
        "keySkills": ["python", "tensorflow", "pytorch", "machine learning"]
    }
]

STOP_WORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from",
    "in", "is", "it", "of", "on", "or", "that", "the", "to", "with"
}

MONTH_PATTERN = (
    r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
)
DATE_RANGE_PATTERN = re.compile(
    rf"(?:{MONTH_PATTERN}\s+\d{{4}}(?:\s*(?:-|to|–|—)\s*(?:present|current|now|{MONTH_PATTERN}\s+\d{{4}}|\d{{4}}))?"
    rf"|\d{{4}}\s*(?:-|to|–|—)\s*(?:present|current|now|\d{{4}})|{MONTH_PATTERN}\s+\d{{4}}|\b\d{{4}}\b)",
    flags=re.IGNORECASE
)
PHONE_PATTERN = re.compile(r"(?:\+?\d[\d\s().-]{8,}\d)")
DETAIL_VERB_HINTS = (
    "developed", "built", "created", "implemented", "designed", "managed",
    "worked", "used", "gained", "delivered", "led", "collaborated",
    "analyzed", "improved", "deployed", "integrated", "maintained",
    "optimized", "supported", "participated", "solved"
)


def clean_text(text):
    text = (text or "").lower()

    if nlp is not None:
        doc = nlp(text)
        tokens = [
            token.lemma_
            for token in doc
            if not token.is_stop and token.is_alpha
        ]
        return " ".join(tokens)

    tokens = re.findall(r"[a-zA-Z]+", text)
    return " ".join(token for token in tokens if token not in STOP_WORDS)


def fallback_similarity(resume_text, job_text):
    cleaned_resume = clean_text(resume_text)
    cleaned_job = clean_text(job_text)

    resume_tokens = cleaned_resume.split()
    job_tokens = cleaned_job.split()

    if not resume_tokens or not job_tokens:
        return 0.0

    resume_counts = Counter(resume_tokens)
    job_counts = Counter(job_tokens)
    shared = sum(
        min(resume_counts[word], job_counts[word])
        for word in set(resume_counts) & set(job_counts)
    )
    keyword_score = (2 * shared / (len(resume_tokens) + len(job_tokens))) * 100
    sequence_score = SequenceMatcher(None, cleaned_resume, cleaned_job).ratio() * 100
    return (keyword_score * 0.7) + (sequence_score * 0.3)


def semantic_match(resume_text, job_text):
    if semantic_model is not None and util is not None:
        cleaned_resume = clean_text(resume_text)
        cleaned_job = clean_text(job_text)

        embeddings = semantic_model.encode(
            [cleaned_resume, cleaned_job],
            convert_to_tensor=True
        )
        similarity = util.cos_sim(embeddings[0], embeddings[1])
        return float(similarity) * 100

    return fallback_similarity(resume_text, job_text)


def extract_skills(text):
    text = (text or "").lower()
    found = set()

    for skill in ALL_SKILLS:
        patterns = [
            r"\b" + re.escape(skill) + r"\b",
            r"\b" + re.escape(skill.replace(" ", "")) + r"\b",
            r"\b" + re.escape(skill.replace(" ", "-")) + r"\b",
        ]

        for pattern in patterns:
            if re.search(pattern, text):
                found.add(skill)
                break

    return list(found)


def build_skill_match_score(job_skills, matched_skills):
    if not job_skills:
        return 0
    return round((len(matched_skills) / len(job_skills)) * 100)


def detect_critical_skills(job_desc, job_skills, matched_skills):
    matched_skill_set = set(matched_skills)
    sentences = re.split(r"[.\n]+", job_desc.lower())
    scored = []

    for skill in job_skills:
        score = 1
        for sentence in sentences:
            if skill in sentence:
                score += 1
                if any(hint in sentence for hint in CRITICAL_SKILL_HINTS):
                    score += 2
        scored.append({
            "skill": skill,
            "matched": skill in matched_skill_set,
            "importance_score": score
        })

    scored.sort(key=lambda item: (-item["importance_score"], item["matched"], item["skill"]))
    return scored[: min(6, len(scored))]


def build_skill_category_progress(job_skills, matched_skills):
    matched_skill_set = set(matched_skills)
    progress = []

    for category, category_skills in SKILL_CATEGORIES.items():
        required = sorted(skill for skill in job_skills if skill in category_skills)
        if not required:
            continue

        matched_in_category = sorted(skill for skill in required if skill in matched_skill_set)
        missing_in_category = sorted(skill for skill in required if skill not in matched_skill_set)
        percent = round((len(matched_in_category) / len(required)) * 100) if required else 0

        progress.append({
            "category": category,
            "matched_count": len(matched_in_category),
            "required_count": len(required),
            "percent": percent,
            "matched_skills": matched_in_category,
            "missing_skills": missing_in_category
        })

    return progress


def build_skill_insight(matched_skills, missing_skills, category_progress):
    if matched_skills:
        strongest = max(category_progress, key=lambda item: item["percent"], default=None)
        matched_text = ", ".join(matched_skills[:3])
        if strongest and strongest["matched_skills"]:
            insight = (
                f"The candidate demonstrates the strongest alignment in {strongest['category'].lower()} skills, "
                f"with evidence of {matched_text} in the resume."
            )
        else:
            insight = f"The candidate demonstrates alignment through resume skills such as {matched_text}."
    else:
        insight = "The resume shows limited direct overlap with the skills extracted from the job description."

    if missing_skills:
        insight += f" Important job skills still missing from the resume include {', '.join(missing_skills[:3])}."
    else:
        insight += " No major skill gaps were identified from the extracted job-description skills."

    return insight


def build_skill_dashboard(job_desc, job_skills, matched_skills, missing_skills):
    category_progress = build_skill_category_progress(job_skills, matched_skills)
    return {
        "skill_match_score": build_skill_match_score(job_skills, matched_skills),
        "required_skill_count": len(job_skills),
        "matched_skill_count": len(matched_skills),
        "missing_skill_count": len(missing_skills),
        "critical_skills": detect_critical_skills(job_desc, job_skills, matched_skills),
        "category_progress": category_progress,
        "recommended_skills": missing_skills,
        "ai_skill_insight": build_skill_insight(matched_skills, missing_skills, category_progress)
    }


def read_resume_file(resume_file):
    try:
        if resume_file.name.endswith(".pdf"):
            if PdfReader is None:
                raise ValueError("PDF support is not available on the server.")

            resume_text = ""
            reader = PdfReader(resume_file)
            for page in reader.pages:
                resume_text += (page.extract_text() or "") + "\n"
            return resume_text

        if resume_file.name.endswith(".docx"):
            if docx is None:
                raise ValueError("DOCX support is not available on the server.")

            document = docx.Document(resume_file)
            return "\n".join(paragraph.text for paragraph in document.paragraphs)

        return resume_file.read().decode(errors="ignore")
    except Exception as exc:
        raise ValueError(str(exc)) from exc


def infer_candidate_profile(resume_text, resume_skills):
    role_signals = [
        ("Machine Learning Engineer", {"machine learning", "tensorflow", "pytorch"}),
        ("Data Analyst", {"sql", "pandas", "numpy", "data science"}),
        ("Backend Developer", {"python", "django", "flask", "fastapi", "api", "sql"}),
        ("Frontend Developer", {"html", "css", "javascript", "react", "angular", "vue"}),
        ("Full Stack Developer", {"python", "django", "javascript", "html", "css", "react"}),
        ("DevOps Engineer", {"docker", "kubernetes", "aws", "azure", "gcp", "ci/cd", "linux"}),
    ]

    resume_skill_set = set(resume_skills)
    best_role = None
    best_score = 0

    for role, signals in role_signals:
        score = len(resume_skill_set & signals)
        if score > best_score:
            best_role = role
            best_score = score

    if best_role:
        return best_role

    lower_resume = resume_text.lower()
    if "developer" in lower_resume:
        return "Developer"
    if "analyst" in lower_resume:
        return "Analyst"
    if "engineer" in lower_resume:
        return "Engineer"
    return "Technical Candidate"


def format_skill_list(skills, limit=5):
    if not skills:
        return ""

    ordered_skills = sorted(skills)
    if len(ordered_skills) <= limit:
        return ", ".join(ordered_skills)

    visible_skills = ordered_skills[:limit]
    return f"{', '.join(visible_skills)}, and {len(ordered_skills) - limit} more"


def list_present_sections(parsed_resume):
    ordered_sections = [
        ("summary", "summary"),
        ("skills", "skills"),
        ("experience", "experience"),
        ("projects", "projects"),
        ("education", "education"),
        ("certifications", "certifications")
    ]

    present = [label for key, label in ordered_sections if parsed_resume["sections"].get(key)]
    if present:
        return present

    if parsed_resume["intro_lines"]:
        return ["header content"]

    return []


def build_summary_sections(resume_text, resume_skills, job_skills, matched_skills, missing_skills):
    candidate_profile = infer_candidate_profile(resume_text, resume_skills)
    parsed_resume = parse_resume_sections(resume_text)
    present_sections = list_present_sections(parsed_resume)
    project_count = len(build_structured_entries(parsed_resume["sections"]["projects"], "projects"))
    experience_count = len(build_structured_entries(parsed_resume["sections"]["experience"], "experience"))
    matched_count = len(matched_skills)
    missing_count = len(missing_skills)
    total_required = len(job_skills)
    matched_skill_text = format_skill_list(matched_skills, limit=4)
    missing_skill_text = format_skill_list(missing_skills, limit=4)
    resume_skill_text = format_skill_list(resume_skills, limit=4)
    section_text = ", ".join(present_sections[:3]) if present_sections else "limited section structure"

    if matched_skills:
        relevant_skills = (
            f"Direct overlap with the job description: {matched_skill_text}. "
            f"That covers {matched_count} of {total_required} extracted required skills."
        )
    elif job_skills:
        relevant_skills = (
            f"No direct overlap was detected with the extracted job-description skills. "
            f"The resume still includes technical terms such as {resume_skill_text or 'general technical content'}."
        )
    else:
        relevant_skills = (
            f"The resume includes technical content such as {resume_skill_text or 'general technical details'}, "
            "but the job description did not provide enough extractable skill terms for a reliable overlap summary."
        )

    if matched_skills:
        insight = (
            f"The resume shows the strongest fit for a {candidate_profile} direction. "
            f"Alignment is supported by sections such as {section_text}, where relevant technologies and role-related work are visible."
        )
    else:
        insight = (
            f"The resume currently reads more like a {candidate_profile} profile, but direct alignment with the extracted job-description skills is limited. "
            f"Relevant evidence may need to be surfaced more clearly inside {section_text}."
        )

    if missing_skills:
        skill_gaps = (
            f"The job description also includes {missing_skill_text}, which was not found in the resume text. "
            "These should be added only if the candidate has real project or work evidence for them."
        )
    elif job_skills:
        skill_gaps = (
            "No major missing skills were identified from the extracted job-description skills."
        )
    else:
        skill_gaps = (
            "No specific job-description skills were extracted for comparison, so the system could not produce a reliable skill-gap list."
        )

    profile_line = (
        f"The resume aligns most closely with a {candidate_profile} profile based on extracted skills such as "
        f"{matched_skill_text or resume_skill_text or 'the technical terms found in the resume'}. "
        f"The document includes usable evidence in {section_text}."
    )

    evidence_snapshot = (
        f"{matched_count} matched skill{'s' if matched_count != 1 else ''}, "
        f"{missing_count} missing skill{'s' if missing_count != 1 else ''}, "
        f"{project_count} project entr{'y' if project_count == 1 else 'ies'} and "
        f"{experience_count} experience entr{'y' if experience_count == 1 else 'ies'} detected."
    )

    sections = {
        "candidate_profile": profile_line,
        "relevant_skills": relevant_skills,
        "job_match_insight": insight,
        "skill_gaps": skill_gaps,
        "analysis_evidence": evidence_snapshot
    }

    summary = (
        "Key Summary\n\n"
        f"Candidate Profile:\n{sections['candidate_profile']}\n\n"
        f"Relevant Skills:\n{sections['relevant_skills']}\n\n"
        f"Job Match Insight:\n{sections['job_match_insight']}\n\n"
        f"Skill Gaps:\n{sections['skill_gaps']}\n\n"
        f"Evidence Snapshot:\n{sections['analysis_evidence']}"
    )

    return sections, summary


def format_section_name(section_key):
    mapping = {
        "summary": "Summary",
        "skills": "Skills",
        "experience": "Experience",
        "projects": "Projects",
        "education": "Education",
        "certifications": "Certifications"
    }
    return mapping.get(section_key, section_key.title())


def find_skill_evidence_sections(skill, parsed_resume):
    matched_sections = []
    skill_pattern = re.compile(r"\b" + re.escape(skill) + r"\b", re.IGNORECASE)

    for section_name, lines in parsed_resume["sections"].items():
        combined_text = " ".join(lines)
        if combined_text and skill_pattern.search(combined_text):
            matched_sections.append(format_section_name(section_name))

    header_text = " ".join(parsed_resume["intro_lines"] + parsed_resume["header_lines"])
    if header_text and skill_pattern.search(header_text):
        matched_sections.append("Header")

    return dedupe_preserve(matched_sections)


def extract_job_signal(skill, job_desc):
    skill_pattern = re.compile(r"\b" + re.escape(skill) + r"\b", re.IGNORECASE)
    sentences = re.split(r"(?<=[.!?])\s+|\n+", job_desc or "")

    for sentence in sentences:
        cleaned_sentence = re.sub(r"\s+", " ", sentence).strip()
        if cleaned_sentence and skill_pattern.search(cleaned_sentence):
            if len(cleaned_sentence) > 180:
                return cleaned_sentence[:177].rstrip() + "..."
            return cleaned_sentence

    return ""


def suggest_resume_section(parsed_resume):
    if parsed_resume["sections"].get("projects"):
        return "Projects"
    if parsed_resume["sections"].get("experience"):
        return "Experience"
    if parsed_resume["sections"].get("skills"):
        return "Skills"
    return "Skills"


def generate_roadmap(parsed_resume, job_desc, missing_skills):
    roadmap = []

    for skill in missing_skills[:5]:
        suggested_section = suggest_resume_section(parsed_resume)
        roadmap.append({
            "skill": skill,
            "category": get_skill_category(skill),
            "phase": "Priority gap from the job description",
            "gap_reason": f"'{skill}' appears in the job description but was not detected in the resume text.",
            "job_signal": extract_job_signal(skill, job_desc),
            "suggested_section": suggested_section,
            "weekly_plan": [
                f"Confirm whether you already have real coursework, project, internship, or work evidence for {skill}.",
                f"If real evidence exists, add a concrete bullet under {suggested_section} that shows how {skill} was used.",
                f"If no direct evidence exists yet, treat {skill} as a genuine learning target before adding it to the resume."
            ]
        })

    return roadmap


def generate_strengths(parsed_resume, matched_skills):
    strength_cards = []

    for skill in matched_skills[:5]:
        evidence_sections = find_skill_evidence_sections(skill, parsed_resume)
        strength_cards.append({
            "skill": skill,
            "category": get_skill_category(skill),
            "evidence_sections": evidence_sections,
            "summary": f"'{skill}' appears in both the resume and the job description.",
            "evidence_note": (
                f"Resume evidence was detected in {', '.join(evidence_sections)}."
                if evidence_sections else
                "The skill was detected in resume text, but no single section could be isolated."
            )
        })

    return strength_cards


def generate_improvements(parsed_resume, job_desc, missing_skills):
    improvement_cards = []

    for skill in missing_skills[:5]:
        suggested_section = suggest_resume_section(parsed_resume)
        improvement_cards.append({
            "skill": skill,
            "category": get_skill_category(skill),
            "job_signal": extract_job_signal(skill, job_desc),
            "reason": f"'{skill}' was found in the job description but not in the resume text.",
            "resume_action": (
                f"Add {skill} under {suggested_section} only if the candidate has real evidence from projects, experience, coursework, or internships."
            ),
            "upskill_action": (
                f"If no evidence exists yet, keep {skill} as an upskilling target instead of claiming it on the resume."
            ),
            "suggested_section": suggested_section
        })

    return improvement_cards


def parse_experience(text):
    matches = re.findall(r"(\d{1,2})(\+)?\s*(?:years|yrs|year)", text.lower())
    if not matches:
        return 0

    return max(int(match[0]) for match in matches)


def parse_education(text):
    education_levels = [
        ("phd", 4),
        ("doctor", 4),
        ("master", 3),
        ("bachelor", 2),
        ("b.sc", 2),
        ("m.sc", 3),
        ("associate", 1)
    ]

    found = []
    lower_text = text.lower()
    for key, value in education_levels:
        if key in lower_text:
            found.append((value, key))

    if not found:
        return None

    found.sort(reverse=True)
    return found[0][1]


def categorize_skills(matched, missing):
    categories = {}
    for category, skills in SKILLS.items():
        matched_in_category = [skill for skill in matched if skill in skills]
        missing_in_category = [skill for skill in missing if skill in skills]
        if matched_in_category or missing_in_category:
            categories[category] = {
                "matched": matched_in_category,
                "missing": missing_in_category
            }
    return categories


def get_skill_category(skill):
    for category, skills in SKILL_CATEGORIES.items():
        if skill in skills:
            return category
    return "Other"


def dedupe_preserve(items):
    seen = set()
    ordered = []

    for item in items:
        cleaned = re.sub(r"\s+", " ", str(item)).strip(" :\t")
        if not cleaned:
            continue

        lookup_key = cleaned.lower()
        if lookup_key in seen:
            continue

        seen.add(lookup_key)
        ordered.append(cleaned)

    return ordered


def split_resume_lines(text):
    normalized_text = (text or "").replace("\u2022", "\n- ").replace("\r", "\n")
    raw_lines = re.split(r"\n+", normalized_text)
    return dedupe_preserve(raw_lines)


def normalize_heading(line):
    lowered = str(line).lower().replace("&", " and ")
    lowered = re.sub(r"[^a-z0-9\s]", " ", lowered)
    return re.sub(r"\s+", " ", lowered).strip()


def detect_section_heading(line):
    heading = normalize_heading(line)
    if not heading:
        return None

    for section, aliases in SECTION_ALIASES.items():
        if heading in aliases:
            return section

    return None


def extract_contact_details(text):
    source_text = text or ""

    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", source_text)
    phone_match = PHONE_PATTERN.search(source_text)
    links = re.findall(
        r"(?:https?://|www\.)\S+|(?:linkedin\.com/\S+)|(?:github\.com/\S+)",
        source_text,
        flags=re.IGNORECASE
    )

    linkedin = next((link for link in links if "linkedin" in link.lower()), "")
    github = next((link for link in links if "github" in link.lower()), "")
    portfolio = next(
        (link for link in links if "linkedin" not in link.lower() and "github" not in link.lower()),
        ""
    )

    return {
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio
    }


def is_contact_line(line):
    lowered = str(line).lower()
    return any([
        re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", line),
        PHONE_PATTERN.search(line),
        "linkedin" in lowered,
        "github" in lowered,
        "portfolio" in lowered,
        "www." in lowered,
        "http://" in lowered,
        "https://" in lowered
    ])


def format_display_name(name):
    words = []
    for word in str(name).split():
        if word.isupper() and len(word) <= 3:
            words.append(word)
        elif word.isupper():
            words.append(word.title())
        else:
            words.append(word)
    return " ".join(words)


def extract_candidate_name(lines):
    for line in lines[:8]:
        normalized = normalize_heading(line)
        words = normalized.split()

        if not words or len(words) < 2 or len(words) > 5:
            continue
        if detect_section_heading(line) or is_contact_line(line):
            continue
        if any(char.isdigit() for char in line):
            continue
        if any(token in {"resume", "curriculum", "vitae"} for token in words):
            continue

        alphabetic_ratio = sum(char.isalpha() for char in line) / max(1, len(line))
        if alphabetic_ratio < 0.6:
            continue

        return format_display_name(line.strip())

    return ""


def parse_resume_sections(text):
    lines = split_resume_lines(text)
    sections = {section: [] for section in SECTION_ALIASES}
    header_lines = []
    current_section = None

    for line in lines:
        heading = detect_section_heading(line)
        if heading:
            current_section = heading
            continue

        if current_section:
            sections[current_section].append(line)
        else:
            header_lines.append(line)

    name = extract_candidate_name(header_lines)
    intro_lines = [
        line for line in header_lines
        if line != name and not is_contact_line(line) and not detect_section_heading(line)
    ]

    return {
        "name": name,
        "header_lines": header_lines,
        "intro_lines": dedupe_preserve(intro_lines),
        "sections": {key: dedupe_preserve(value) for key, value in sections.items()}
    }


def build_ats_skill_lines(resume_skills):
    grouped_skills = {}
    for skill in sorted(resume_skills):
        grouped_skills.setdefault(get_skill_category(skill), []).append(skill)

    ordered_categories = ["Programming", "Web Development", "Database", "Cloud", "Tools", "Other"]
    lines = []

    for category in ordered_categories:
        category_skills = grouped_skills.get(category, [])
        if category_skills:
            lines.append(f"{category}: {', '.join(category_skills)}")

    return lines


def append_ats_section(output_lines, heading, values, bullet_style=False):
    if not values:
        return

    output_lines.append("")
    output_lines.append(heading)
    for value in values:
        if bullet_style:
            output_lines.append(f"- {value}")
        else:
            output_lines.append(value)


def strip_bullet_prefix(value):
    return re.sub(r"^[\-\u2022*]+\s*", "", str(value)).strip()


def split_line_with_meta(value):
    cleaned = strip_bullet_prefix(value)
    date_match = DATE_RANGE_PATTERN.search(cleaned)

    if date_match and date_match.start() >= max(10, len(cleaned) // 3):
        main_text = cleaned[:date_match.start()].strip(" |-,:")
        meta_text = cleaned[date_match.start():].strip(" |-,:")
        if main_text:
            return main_text, meta_text

    if cleaned.count(",") >= 2 and ")" in cleaned:
        bracket_index = cleaned.rfind(")")
        if bracket_index != -1 and bracket_index < len(cleaned) - 2:
            main_text = cleaned[: bracket_index + 1].strip(" |-,:")
            meta_text = cleaned[bracket_index + 1 :].strip(" |-,:")
            if main_text and meta_text:
                return main_text, meta_text

    return cleaned, ""


def looks_like_metadata_line(value):
    cleaned = strip_bullet_prefix(value)
    if not cleaned:
        return False

    lowered = cleaned.lower()
    if re.fullmatch(DATE_RANGE_PATTERN, cleaned):
        return True
    if DATE_RANGE_PATTERN.search(cleaned) and len(cleaned.split()) <= 10:
        return True
    if cleaned.count(",") >= 2 and len(cleaned.split()) <= 14 and not cleaned.endswith("."):
        return True
    if any(token in lowered for token in ("remote", "india", "hyderabad", "vijayawada", "bangalore", "location")):
        return len(cleaned.split()) <= 14

    return False


def looks_like_detail_line(value):
    original = str(value).strip()
    cleaned = strip_bullet_prefix(value)
    lowered = cleaned.lower()

    return any([
        original.startswith(("-", "\u2022", "*")),
        cleaned.endswith("."),
        len(cleaned.split()) > 18,
        lowered.startswith(DETAIL_VERB_HINTS),
        cleaned.startswith("http://"),
        cleaned.startswith("https://"),
        cleaned.startswith("www."),
    ])


def should_start_new_entry(value, current_entry, section_name):
    cleaned = strip_bullet_prefix(value)

    if not current_entry or not current_entry.get("title"):
        return False
    if looks_like_detail_line(cleaned) or looks_like_metadata_line(cleaned):
        return False
    if current_entry["details"]:
        return True
    if section_name in {"education", "certifications"} and current_entry["subtitle"]:
        return True
    if section_name in {"experience", "projects"} and current_entry["subtitle"] and len(cleaned.split()) <= 14:
        return True

    return False


def build_structured_entries(lines, section_name):
    entries = []
    current_entry = None

    for line in lines:
        raw_text = str(line).strip()
        if not raw_text:
            continue

        cleaned_text = strip_bullet_prefix(raw_text)
        main_text, inline_meta = split_line_with_meta(raw_text)

        if current_entry is None:
            current_entry = {
                "title": main_text,
                "title_meta": inline_meta,
                "subtitle": "",
                "subtitle_meta": "",
                "details": []
            }
            continue

        if looks_like_detail_line(raw_text):
            current_entry["details"].append(cleaned_text)
            continue

        if looks_like_metadata_line(cleaned_text):
            if not current_entry["title_meta"]:
                current_entry["title_meta"] = cleaned_text
            elif not current_entry["subtitle_meta"]:
                current_entry["subtitle_meta"] = cleaned_text
            else:
                current_entry["details"].append(cleaned_text)
            continue

        if should_start_new_entry(cleaned_text, current_entry, section_name):
            entries.append(current_entry)
            current_entry = {
                "title": main_text,
                "title_meta": inline_meta,
                "subtitle": "",
                "subtitle_meta": "",
                "details": []
            }
            continue

        if not current_entry["subtitle"]:
            current_entry["subtitle"] = main_text
            if inline_meta:
                current_entry["subtitle_meta"] = inline_meta
            continue

        if inline_meta and not current_entry["subtitle_meta"]:
            current_entry["subtitle_meta"] = inline_meta
            if main_text and main_text != cleaned_text:
                current_entry["details"].append(main_text)
            elif main_text and main_text != current_entry["subtitle"]:
                current_entry["details"].append(main_text)
            continue

        current_entry["details"].append(cleaned_text)

    if current_entry:
        entries.append(current_entry)

    return entries


def append_structured_entries(output_lines, heading, entries):
    if not entries:
        return

    output_lines.append("")
    output_lines.append(heading)

    for entry in entries:
        title_line = entry["title"]
        if entry["title_meta"]:
            title_line = f"{title_line} | {entry['title_meta']}"
        output_lines.append(title_line)

        if entry["subtitle"]:
            subtitle_line = entry["subtitle"]
            if entry["subtitle_meta"]:
                subtitle_line = f"{subtitle_line} | {entry['subtitle_meta']}"
            output_lines.append(subtitle_line)

        for detail in entry["details"]:
            output_lines.append(f"- {detail}")


def build_ats_resume_text(parsed_resume, contacts, resume_skills):
    sections = parsed_resume["sections"]
    summary_lines = sections["summary"] or parsed_resume["intro_lines"][:3]
    skill_lines = sections["skills"] or build_ats_skill_lines(resume_skills)
    education_entries = build_structured_entries(sections["education"], "education")
    experience_entries = build_structured_entries(sections["experience"], "experience")
    project_entries = build_structured_entries(sections["projects"], "projects")
    certification_entries = build_structured_entries(sections["certifications"], "certifications")

    contact_parts = []
    if contacts["email"]:
        contact_parts.append(contacts["email"])
    if contacts["phone"]:
        contact_parts.append(contacts["phone"])
    for key in ("linkedin", "github", "portfolio"):
        if contacts[key]:
            contact_parts.append(contacts[key])

    output_lines = [parsed_resume["name"].upper() if parsed_resume["name"] else "CANDIDATE"]
    if contact_parts:
        output_lines.append(" | ".join(contact_parts))

    append_ats_section(output_lines, "SUMMARY", summary_lines, bullet_style=True)
    append_structured_entries(output_lines, "EDUCATION", education_entries)
    append_structured_entries(output_lines, "EXPERIENCE", experience_entries)
    append_structured_entries(output_lines, "PROJECTS", project_entries)
    append_ats_section(output_lines, "SKILLS", skill_lines)
    append_structured_entries(output_lines, "CERTIFICATIONS", certification_entries)

    return "\n".join(output_lines).strip()


def build_ats_analysis(resume_text, resume_file_name, resume_skills, job_skills, matched_skills, missing_skills):
    parsed_resume = parse_resume_sections(resume_text)
    contacts = extract_contact_details(resume_text)
    sections = parsed_resume["sections"]

    summary_lines = sections["summary"] or parsed_resume["intro_lines"][:3]
    skill_lines = sections["skills"] or build_ats_skill_lines(resume_skills)
    has_link = any([contacts["linkedin"], contacts["github"], contacts["portfolio"]])
    has_skills_section = bool(sections["skills"])
    has_skills_content = bool(has_skills_section or resume_skills)
    has_experience_or_projects = bool(sections["experience"] or sections["projects"])
    has_education = bool(sections["education"])
    detail_line_count = len(sections["experience"]) + len(sections["projects"])
    keyword_coverage = build_skill_match_score(job_skills, matched_skills) if job_skills else 0

    checks = []

    def add_check(label, status, awarded, possible, detail):
        checks.append({
            "label": label,
            "status": status,
            "points_awarded": awarded,
            "points_possible": possible,
            "detail": detail
        })

    add_check(
        "Email Address",
        "passed" if contacts["email"] else "missing",
        8 if contacts["email"] else 0,
        8,
        "Professional email detected in the resume." if contacts["email"] else "No email address was detected."
    )
    add_check(
        "Phone Number",
        "passed" if contacts["phone"] else "missing",
        8 if contacts["phone"] else 0,
        8,
        "Phone number detected in the resume." if contacts["phone"] else "No phone number was detected."
    )
    add_check(
        "Professional Links",
        "passed" if has_link else "missing",
        4 if has_link else 0,
        4,
        "LinkedIn, GitHub, or portfolio link detected." if has_link else "No LinkedIn, GitHub, or portfolio link was detected."
    )
    add_check(
        "Summary Section",
        "passed" if summary_lines else "missing",
        5 if summary_lines else 0,
        5,
        "Summary or profile text was detected near the top of the resume." if summary_lines else "No summary or objective section was detected."
    )

    skills_points = 10 if has_skills_section else 6 if has_skills_content else 0
    skills_status = "passed" if has_skills_section else "partial" if has_skills_content else "missing"
    add_check(
        "Skills Section",
        skills_status,
        skills_points,
        10,
        "A dedicated skills section was detected." if has_skills_section else (
            "Skills were detected in the resume text, but not under a clear skills heading." if has_skills_content
            else "No clear technical skills section or extracted skills were detected."
        )
    )

    add_check(
        "Experience or Projects",
        "passed" if has_experience_or_projects else "missing",
        10 if has_experience_or_projects else 0,
        10,
        "Experience or projects content was detected." if has_experience_or_projects else "No experience or projects section was clearly detected."
    )
    add_check(
        "Education Section",
        "passed" if has_education else "missing",
        10 if has_education else 0,
        10,
        "Education details were detected." if has_education else "No education section was clearly detected."
    )

    keyword_points = round(keyword_coverage * 0.35)
    keyword_status = "passed" if keyword_coverage >= 70 else "partial" if keyword_coverage > 0 else "missing"
    add_check(
        "Job Keyword Coverage",
        keyword_status,
        keyword_points,
        35,
        (
            f"{len(matched_skills)} of {len(job_skills)} extracted job skills also appear in the resume."
            if job_skills else
            "No job-description skills were extracted for ATS keyword comparison."
        )
    )

    detail_points = 10 if detail_line_count >= 3 else 5 if detail_line_count else 0
    detail_status = "passed" if detail_line_count >= 3 else "partial" if detail_line_count else "missing"
    add_check(
        "Structured Detail Lines",
        detail_status,
        detail_points,
        10,
        (
            "Multiple experience or project detail lines were detected, which helps ATS readability."
            if detail_line_count >= 3 else
            "Some experience or project details were detected, but the resume could use more structured bullet content."
            if detail_line_count else
            "No structured experience or project detail lines were detected."
        )
    )

    score = sum(item["points_awarded"] for item in checks)
    improvements = []

    if not contacts["email"]:
        improvements.append("Add a professional email address near the top of the resume.")
    if not contacts["phone"]:
        improvements.append("Add a phone number in the resume header for ATS completeness.")
    if not has_link:
        improvements.append("Add a LinkedIn, GitHub, or portfolio link to strengthen the resume header.")
    if not summary_lines:
        improvements.append("Add a short professional summary or objective section at the top of the resume.")
    if not has_skills_section:
        improvements.append("Add a dedicated 'Skills' section so ATS systems can scan technologies more reliably.")
    if not has_experience_or_projects:
        improvements.append("Add an experience or projects section with clear bullet points from your actual work.")
    elif detail_line_count < 3:
        improvements.append("Use more structured bullet points under experience or projects to improve ATS readability.")
    if not has_education:
        improvements.append("Add a clear education section with degree, college, and graduation details.")
    if missing_skills:
        improvements.append(
            f"Important job-description skills not found in the resume: {', '.join(missing_skills[:5])}."
        )

    resume_stub = re.sub(r"[^A-Za-z0-9_-]+", "_", resume_file_name.rsplit(".", 1)[0]).strip("_") or "resume"
    section_payload = {
        "candidate_name": parsed_resume["name"],
        "contacts": contacts,
        "summary_lines": summary_lines,
        "skills_lines": skill_lines,
        "experience_lines": sections["experience"],
        "project_lines": sections["projects"],
        "education_lines": sections["education"],
        "certification_lines": sections["certifications"]
    }

    return {
        "score": score,
        "score_label": (
            "Strong ATS readiness" if score >= 80 else
            "Moderate ATS readiness" if score >= 60 else
            "Needs ATS-focused improvements"
        ),
        "keyword_coverage": keyword_coverage,
        "sections_found": sum(
            1 for present in [bool(summary_lines), has_skills_content, has_experience_or_projects, has_education]
            if present
        ),
        "checks": checks,
        "improvements": improvements,
        "resume_sections": section_payload,
        "ats_resume_text": build_ats_resume_text(parsed_resume, contacts, resume_skills),
        "download_file_name": f"{resume_stub}_ats_resume.docx"
    }


def set_paragraph_bottom_border(paragraph, color="A6A6A6", size="8"):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(docx.oxml.ns.qn("w:pBdr"))

    if borders is None:
        borders = docx.oxml.OxmlElement("w:pBdr")
        paragraph_properties.append(borders)

    bottom = docx.oxml.OxmlElement("w:bottom")
    bottom.set(docx.oxml.ns.qn("w:val"), "single")
    bottom.set(docx.oxml.ns.qn("w:sz"), size)
    bottom.set(docx.oxml.ns.qn("w:space"), "1")
    bottom.set(docx.oxml.ns.qn("w:color"), color)
    borders.append(bottom)


def add_section_heading(document, heading):
    heading_paragraph = document.add_paragraph()
    heading_paragraph.paragraph_format.space_before = docx.shared.Pt(10)
    heading_paragraph.paragraph_format.space_after = docx.shared.Pt(4)
    heading_paragraph.paragraph_format.keep_with_next = True
    heading_run = heading_paragraph.add_run(heading)
    heading_run.bold = True
    heading_run.font.name = "Times New Roman"
    heading_run.font.size = docx.shared.Pt(11)
    set_paragraph_bottom_border(heading_paragraph)
    return heading_paragraph


def add_text_with_meta(
    document,
    left_text,
    right_text="",
    *,
    left_bold=False,
    left_underline=False,
    right_italic=True,
    font_size=11,
    keep_with_next=False
):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = docx.shared.Pt(0)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        docx.shared.Inches(6.9),
        docx.enum.text.WD_TAB_ALIGNMENT.RIGHT
    )

    left_run = paragraph.add_run(left_text)
    left_run.bold = left_bold
    left_run.underline = left_underline
    left_run.font.name = "Times New Roman"
    left_run.font.size = docx.shared.Pt(font_size)

    if right_text:
        right_run = paragraph.add_run(f"\t{right_text}")
        right_run.italic = right_italic
        right_run.font.name = "Times New Roman"
        right_run.font.size = docx.shared.Pt(max(10, font_size - 1))

    return paragraph


def add_summary_to_document(document, values):
    if not values:
        return

    add_section_heading(document, "SUMMARY")
    summary_paragraph = document.add_paragraph()
    summary_paragraph.paragraph_format.space_after = docx.shared.Pt(4)
    summary_paragraph.paragraph_format.keep_together = True
    summary_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    summary_run = summary_paragraph.add_run(" ".join(strip_bullet_prefix(value) for value in values))
    summary_run.font.name = "Times New Roman"
    summary_run.font.size = docx.shared.Pt(10.5)


def add_skills_to_document(document, values):
    if not values:
        return

    add_section_heading(document, "SKILLS")
    for value in values:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = docx.shared.Pt(0)
        paragraph.paragraph_format.keep_together = True
        cleaned = strip_bullet_prefix(value)

        if ":" in cleaned:
            label, detail = cleaned.split(":", 1)
            label_run = paragraph.add_run(f"{label.strip()}: ")
            label_run.bold = True
            label_run.font.name = "Times New Roman"
            label_run.font.size = docx.shared.Pt(10.5)

            detail_run = paragraph.add_run(detail.strip())
            detail_run.font.name = "Times New Roman"
            detail_run.font.size = docx.shared.Pt(10.5)
        else:
            run = paragraph.add_run(cleaned)
            run.font.name = "Times New Roman"
            run.font.size = docx.shared.Pt(10.5)


def add_entry_section_to_document(document, heading, entries):
    if not entries:
        return

    add_section_heading(document, heading)
    for entry in entries:
        add_text_with_meta(
            document,
            entry["title"],
            entry["title_meta"],
            left_bold=True,
            right_italic=True,
            font_size=11,
            keep_with_next=True
        )

        if entry["subtitle"] or entry["subtitle_meta"]:
            subtitle_is_link = entry["subtitle"].startswith(("http://", "https://", "www.")) if entry["subtitle"] else False
            subtitle_paragraph = add_text_with_meta(
                document,
                entry["subtitle"] or "",
                entry["subtitle_meta"],
                left_bold=bool(entry["subtitle"]) and not subtitle_is_link,
                left_underline=subtitle_is_link,
                right_italic=True,
                font_size=10.5,
                keep_with_next=bool(entry["details"])
            )
            if subtitle_is_link:
                for run in subtitle_paragraph.runs:
                    run.underline = True

        for detail in entry["details"]:
            bullet_paragraph = document.add_paragraph(style="List Bullet")
            bullet_paragraph.paragraph_format.space_after = docx.shared.Pt(0)
            bullet_paragraph.paragraph_format.left_indent = docx.shared.Inches(0.18)
            bullet_paragraph.paragraph_format.keep_together = True
            bullet_run = bullet_paragraph.add_run(strip_bullet_prefix(detail))
            bullet_run.font.name = "Times New Roman"
            bullet_run.font.size = docx.shared.Pt(10.25)


def build_ats_resume_docx(ats_analysis):
    if docx is None:
        raise ValueError("DOCX resume download is not available on the server right now.")

    resume_sections = ats_analysis.get("resume_sections") or {}
    contacts = resume_sections.get("contacts") or {}
    candidate_name = (resume_sections.get("candidate_name") or "Candidate").strip()

    document = docx.Document()
    styles = document.styles
    if "Normal" in styles:
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = docx.shared.Pt(10.5)
    if "List Bullet" in styles:
        styles["List Bullet"].font.name = "Times New Roman"
        styles["List Bullet"].font.size = docx.shared.Pt(10.25)

    section = document.sections[0]
    section.top_margin = docx.shared.Inches(0.45)
    section.bottom_margin = docx.shared.Inches(0.45)
    section.left_margin = docx.shared.Inches(0.45)
    section.right_margin = docx.shared.Inches(0.45)

    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_after = docx.shared.Pt(2)
    name_run = name_paragraph.add_run(candidate_name.upper())
    name_run.bold = True
    name_run.font.name = "Times New Roman"
    name_run.font.size = docx.shared.Pt(15.5)

    contact_parts = []
    if contacts.get("email"):
        contact_parts.append(contacts["email"])
    if contacts.get("phone"):
        contact_parts.append(contacts["phone"])
    for key in ("linkedin", "github", "portfolio"):
        if contacts.get(key):
            contact_parts.append(contacts[key])

    if contact_parts:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.paragraph_format.space_after = docx.shared.Pt(3)

        for index, part in enumerate(contact_parts):
            contact_run = contact_paragraph.add_run(part)
            contact_run.font.name = "Times New Roman"
            contact_run.font.size = docx.shared.Pt(9.5)
            contact_run.underline = True
            if index < len(contact_parts) - 1:
                spacer = contact_paragraph.add_run("    ")
                spacer.font.name = "Times New Roman"
                spacer.font.size = docx.shared.Pt(9.5)

    add_summary_to_document(document, resume_sections.get("summary_lines") or [])
    add_entry_section_to_document(
        document,
        "EDUCATION",
        build_structured_entries(resume_sections.get("education_lines") or [], "education")
    )
    add_entry_section_to_document(
        document,
        "EXPERIENCE",
        build_structured_entries(resume_sections.get("experience_lines") or [], "experience")
    )
    add_entry_section_to_document(
        document,
        "PROJECTS",
        build_structured_entries(resume_sections.get("project_lines") or [], "projects")
    )
    add_skills_to_document(document, resume_sections.get("skills_lines") or [])
    add_entry_section_to_document(
        document,
        "CERTIFICATIONS",
        build_structured_entries(resume_sections.get("certification_lines") or [], "certifications")
    )

    file_buffer = BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)
    return file_buffer.getvalue()


def home(request):
    return render(request, "analyzer/index.html")


def upload_page(request):
    return render(request, "analyzer/upload.html")


def result_page(request):
    return render(request, "analyzer/result.html")


@csrf_exempt
def download_ats_resume(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST method allowed"}, status=405)

    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON payload"}, status=400)

    ats_analysis = payload.get("ats_analysis") or {}
    if not ats_analysis.get("resume_sections"):
        return JsonResponse({"error": "ATS resume data is missing. Please analyze the resume again."}, status=400)

    try:
        docx_content = build_ats_resume_docx(ats_analysis)
    except ValueError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except Exception as exc:
        return JsonResponse({"error": str(exc)}, status=500)

    download_name = ats_analysis.get("download_file_name") or "ats_resume.docx"
    response = HttpResponse(
        docx_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{download_name}"'
    return response


@csrf_exempt
def analyze_api(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST method allowed"}, status=405)

    try:
        resume_file = request.FILES.get("resume")
        job_desc = request.POST.get("job_desc", "").strip()

        if not resume_file:
            return JsonResponse({"error": "Resume file is required"}, status=400)

        if not job_desc:
            return JsonResponse({"error": "Job description is required"}, status=400)

        if len(job_desc) < 50:
            return JsonResponse({"error": "Job description must be at least 50 characters"}, status=400)

        if resume_file.size > 10 * 1024 * 1024:
            return JsonResponse({"error": "File size must be less than 10MB"}, status=400)

        allowed_extensions = ["pdf", "doc", "docx", "txt"]
        file_extension = resume_file.name.split(".")[-1].lower()
        if file_extension not in allowed_extensions:
            return JsonResponse(
                {"error": f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"},
                status=400
            )

        resume_text = read_resume_file(resume_file)
        if not resume_text or len(resume_text.strip()) < 100:
            return JsonResponse({"error": "Resume content appears to be too short"}, status=400)

        exp_years = parse_experience(resume_text)
        req_exp = parse_experience(job_desc)
        education = parse_education(resume_text)

        semantic_score = semantic_match(resume_text, job_desc)
        resume_skills = extract_skills(resume_text)
        job_skills = extract_skills(job_desc)

        matched_skills = list(set(resume_skills) & set(job_skills))
        missing_skills = list(set(job_skills) - set(resume_skills))
        skill_dashboard = build_skill_dashboard(job_desc, job_skills, matched_skills, missing_skills)
        ats_analysis = build_ats_analysis(
            resume_text=resume_text,
            resume_file_name=resume_file.name,
            resume_skills=resume_skills,
            job_skills=job_skills,
            matched_skills=matched_skills,
            missing_skills=missing_skills
        )

        skill_overlap_score = (len(matched_skills) / max(1, len(job_skills)) * 100)
        match_percent = (
            semantic_score * 0.6 +
            skill_overlap_score * 0.4
        )

        job_matches = []
        for role in JOB_ROLES:
            total = len(role["keySkills"])
            matched = sum(1 for skill in role["keySkills"] if skill in resume_skills)
            score = round((matched / total) * 100)
            job_matches.append({
                "title": role["title"],
                "description": role["description"],
                "score": score,
                "matched": matched,
                "total": total,
                "missing": [skill for skill in role["keySkills"] if skill not in resume_skills]
            })

        job_matches.sort(key=lambda job: job["score"], reverse=True)

        parsed_resume = parse_resume_sections(resume_text)

        summary_sections, summary = build_summary_sections(
            resume_text=resume_text,
            resume_skills=resume_skills,
            job_skills=job_skills,
            matched_skills=matched_skills,
            missing_skills=missing_skills
        )
        roadmap = generate_roadmap(parsed_resume, job_desc, missing_skills)
        strengths = generate_strengths(parsed_resume, matched_skills)
        improvements = generate_improvements(parsed_resume, job_desc, missing_skills)

        return JsonResponse({
            "summary": summary,
            "match_percent": round(match_percent, 2),
            "experience_years": exp_years,
            "required_experience": req_exp,
            "education": education,
            "matched_skills": matched_skills,
            "missing_skills": missing_skills,
            "job_matches": job_matches,
            "roadmap": roadmap,
            "strengths": strengths,
            "improvements": improvements,
            "skill_categories": categorize_skills(matched_skills, missing_skills),
            "analysis_mode": "ai" if semantic_model is not None else "fallback",
            "summary_sections": summary_sections,
            "skill_dashboard": skill_dashboard,
            "ats_analysis": ats_analysis
        })
    except ValueError as exc:
        return JsonResponse({"error": str(exc)}, status=400)
    except Exception as exc:
        print(f"Analysis error: {str(exc)}")
        return JsonResponse({"error": str(exc)}, status=500)
